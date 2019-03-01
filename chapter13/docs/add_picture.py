import docx

doc = docx.Document()

doc.add_picture('zophie.png', width=docx.shared.Inches(1))

doc.save('add_picture.docx')