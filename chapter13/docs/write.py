import docx

doc = docx.Document()
doc.add_paragraph('Hello world!', 'Title')

paraObj1 = doc.add_paragraph('second para')
paraObj2 = doc.add_paragraph('another para')
paraObj1.add_run(' This text is being added to the second paragraph.')
doc.save('multipleParagraphs.docx')