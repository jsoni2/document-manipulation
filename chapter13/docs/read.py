import docx

doc = docx.Document('demo.docx')

print(len(doc.paragraphs))

print(doc.paragraphs[0].text)

print(len(doc.paragraphs[0].runs))

print(doc.paragraphs[1].text)

print(len(doc.paragraphs[1].runs))

for para in doc.paragraphs[1].runs:
    print(para.text)